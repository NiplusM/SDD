from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("artifacts/Code_Review_V1_Main_User_Flow_RU_v11.docx")

INK = RGBColor(31, 45, 61)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(94, 104, 116)
DECISION = RGBColor(145, 86, 0)
BULLET_NUM_ID = None


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text, url, color="2E74B5", size=11):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), "Calibri")
    run_properties.append(fonts)

    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(font_size)

    font_size_complex = OxmlElement("w:szCs")
    font_size_complex.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(font_size_complex)

    link_color = OxmlElement("w:color")
    link_color.set(qn("w:val"), color)
    run_properties.append(link_color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")

    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_bottom_border(paragraph, color="D7E3F0", size="10"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    set_run_font(run, size=9, color=MUTED)


def create_bullet_numbering(doc):
    """Create a renderer-safe native Word bullet list using a normal Unicode bullet."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId"), "").isdigit()
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId"), "").isdigit()
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•")
    level.append(text)
    justify = OxmlElement("w:lvlJc")
    justify.set(qn("w:val"), "left")
    level.append(justify)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Arial")
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    r_pr.append(size)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_ref)


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    apply_numbering(p, BULLET_NUM_ID)
    # Explicit geometry prevents Word and LibreOffice from shifting markers or
    # wrapped lines when adjacent headings use manual labels such as 1A/5B.
    p.paragraph_format.left_indent = Inches(0.50)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(0.50), WD_TAB_ALIGNMENT.LEFT
    )
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=11, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK)
    return p


def add_decision_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.08

    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "FFF4D6")
    p_pr.append(shading)

    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), "D89A2B")
    borders.append(left)
    p_pr.append(borders)

    label = p.add_run("ТРЕБУЕТ ПРОДУКТОВОГО РЕШЕНИЯ · ")
    set_run_font(label, size=8.5, color=DECISION, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=10.5, color=INK)
    return p


def add_info_note(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.08

    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "EEF5FC")
    p_pr.append(shading)

    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    p_pr.append(borders)

    label_run = p.add_run(f"{label.upper()} · ")
    set_run_font(label_run, size=8.5, color=BLUE, bold=True)
    body_run = p.add_run(text)
    set_run_font(body_run, size=10.5, color=INK)
    return p


def add_scenario_card(doc, number, title, intent, path, result):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.keep_together = True

    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F4F7FA")
    p_pr.append(shading)

    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), "2E74B5")
    borders.append(left)
    p_pr.append(borders)

    label = p.add_run(f"СЦЕНАРИЙ {number} · {title}\n")
    set_run_font(label, size=10.5, color=BLUE, bold=True)
    intent_label = p.add_run("Задача пользователя: ")
    set_run_font(intent_label, size=10.5, color=INK, bold=True)
    intent_run = p.add_run(f"{intent}\n")
    set_run_font(intent_run, size=10.5, color=INK)
    path_label = p.add_run("Путь: ")
    set_run_font(path_label, size=10.5, color=INK, bold=True)
    path_run = p.add_run(f"{path}\n")
    set_run_font(path_run, size=10.5, color=INK)
    result_label = p.add_run("Результат: ")
    set_run_font(result_label, size=10.5, color=INK, bold=True)
    result_run = p.add_run(result)
    set_run_font(result_run, size=10.5, color=INK)
    return p


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_status_table(doc):
    rows = [
        (
            "Submit Review",
            "Продолжает текущее ревью: позволяет добавить сообщение агенту, собирает решения, исправления и комментарии и запускает новую итерацию.",
            "Open → Updating → Updated",
        ),
        (
            "Complete Review",
            "Завершает активное review, сохраняет итоговое состояние findings и переводит результат в режим только для чтения.",
            "Completed",
        ),
        (
            "Cancel Review",
            "Отменяет само ревью и предложенные в Preview изменения. Это отдельный отказ от результата, а не альтернатива штатному завершению.",
            "Cancelled",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1750, 5050, 2440])
    table.allow_autofit = False
    set_repeat_table_header(table.rows[0])

    headers = ("Действие", "Что происходит", "Статус review")
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = 1
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "DCEAF7")
        tc_pr.append(shd)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=10, color=DARK_BLUE, bold=True)

    for action, behavior, status in rows:
        cells = table.add_row().cells
        for idx, value in enumerate((action, behavior, status)):
            cells[idx].vertical_alignment = 1
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(value)
            set_run_font(run, size=9.8, color=INK, bold=(idx == 0))

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    return table


def add_labeled_heading(doc, style_name, label, title, text_position):
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.left_indent = Inches(text_position)
    p.paragraph_format.first_line_indent = Inches(-text_position)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(text_position), WD_TAB_ALIGNMENT.LEFT
    )
    label_run = p.add_run(str(label))
    label_size = 13 if style_name == "Heading 2" else 12
    label_color = BLUE if style_name == "Heading 2" else DARK_BLUE
    set_run_font(label_run, size=label_size, color=label_color, bold=True)
    tab_run = p.add_run("\t")
    set_run_font(tab_run, size=label_size, color=label_color, bold=True)
    title_run = p.add_run(title)
    set_run_font(title_run, size=label_size, color=label_color, bold=True)
    return p


def add_substep(doc, label, title, bullets):
    add_labeled_heading(doc, "Heading 3", label, title, 0.46)
    for text in bullets:
        add_bullet(doc, text)


def add_step(doc, number, title, paragraphs=None, bullets=None, page_break_before=False):
    p = add_labeled_heading(doc, "Heading 2", f"{number}.", title, 0.36)
    p.paragraph_format.page_break_before = page_break_before
    for text in paragraphs or []:
        add_body(doc, text)
    for text in bullets or []:
        add_bullet(doc, text)


doc = Document()
BULLET_NUM_ID = create_bullet_numbering(doc)
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.82)
section.bottom_margin = Inches(0.82)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.42)
section.footer_distance = Inches(0.42)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("CODE REVIEW · V1 USER FLOW")
set_run_font(hr, size=8.5, color=MUTED, bold=True)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = fp.add_run("Страница ")
set_run_font(fr, size=9, color=MUTED)
add_field(fp, "PAGE")

kicker = doc.add_paragraph()
kicker.paragraph_format.space_before = Pt(10)
kicker.paragraph_format.space_after = Pt(5)
kr = kicker.add_run("PRODUCT FLOW SPECIFICATION")
set_run_font(kr, size=9, color=BLUE, bold=True)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(5)
tr = title.add_run("Code Review изменений в IDE")
set_run_font(tr, size=24, color=INK, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(12)
sr = subtitle.add_run("Основной пользовательский flow V1: от выбора scope до решения о коммите")
set_run_font(sr, size=12.5, color=MUTED)

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(12)
mr1 = meta.add_run("Статус: ")
set_run_font(mr1, size=10, color=INK, bold=True)
mr2 = meta.add_run("целевая спецификация V1")
set_run_font(mr2, size=10, color=INK)
mr3 = meta.add_run("    Область: ")
set_run_font(mr3, size=10, color=INK, bold=True)
mr4 = meta.add_run("GUI review перед коммитом")
set_run_font(mr4, size=10, color=INK)
add_bottom_border(meta)

h = doc.add_paragraph(style="Heading 1")
h.add_run("Проблема")
add_body(doc, "После завершения задачи или итерации пользователю трудно быстро понять, готовы ли изменения к коммиту: весь ли нужный scope проверен, какие риски действительно важны и какие замечания ещё требуют решения.")
add_body(doc, "По мере исправлений и повторных проверок результаты ревью могут расходиться с актуальным кодом. Без единого состояния пользователь теряет связь между найденной проблемой, принятым решением и итоговым набором незакоммиченных изменений.")

h = doc.add_paragraph(style="Heading 1")
h.add_run("Основной пользовательский flow")

add_step(doc, 1, "Пользователь выбирает, какие изменения проверить")
add_body(doc, "Пользователь закончил задачу или очередную итерацию и хочет проверить изменения перед коммитом. Code Review в V1 является только GUI-функцией IDE: Terminal и CLI не входят в scope версии.")
add_body(doc, "Точка входа определяет исходный scope и способ настройки. Любая точка входа активна только при наличии reviewable changes; если scope пуст, IDE не предлагает запуск и объясняет причину: No changes to review.")

add_substep(doc, "1A.", "Commit tool window — весь подготовленный commit scope", (
    "Пользователь запускает Code Review из списка изменений, подготовленных к коммиту.",
    "В исходный scope попадают включённые modified, added, deleted, renamed и untracked files.",
))

add_substep(doc, "1B.", "Control + Control — контекст активного tool window", (
    "Пользователь открывает общий диалог быстрых действий двойным нажатием Control.",
    "IDE формирует scope из изменений активного tool window или editor context; действие Code Review уже выбрано. Текст запроса описывает проверку всего сформированного scope и сохраняет его рабочий контекст, а не сводится к узкой команде ревью кода.",
    "Если активный контекст не содержит сформированных изменений, действие недоступно.",
))

add_substep(doc, "1C.", "Diff — узкий scope выбранных изменений", (
    "Пользователь запускает ревью из открытого diff, когда хочет проверить не весь набор изменений, а выбранный файл, группу файлов или часть изменений.",
    "Выбранный diff становится исходным scope и сохраняется как source context ревью.",
))

add_substep(doc, "1D.", "Текущий чат — прямой запуск тем же агентом", (
    "По умолчанию scope включает все changes текущего чата. Комментарии из чата и связанных файлов также становятся контекстом проверки.",
    "Пользователь может добавить контекст или сузить scope, прикрепив к сообщению нужные файлы и другие доступные attachments.",
    "Перед запуском он указывает другую модель, Effort и при необходимости дополнительные инструкции. Сменить агента в текущем чате нельзя.",
    "Команда запускает ревью сразу в текущем чате тем же агентом; отдельный launch dialog и новая чат-сессия не создаются.",
))

add_step(doc, 2, "Пользователь настраивает ревью")
add_body(doc, "Для Commit tool window, Control + Control и diff единый launch dialog показывает уже сформированный scope и параметры продолжения работы. Пользователь понимает, что именно будет проверено, и может изменить настройку до запуска.")
for item in (
    "Code Review — предвыбранный тип действия. Связанный с ним текст формулирует полноценную задачу на проверку выбранного scope: какие изменения и контекст нужно учесть, на что обратить внимание и какой результат ожидается.",
    "Scope — весь доступный контекст или выбранная часть; файлы и изменения можно добавить или исключить.",
    "Session — New Session по умолчанию или Current Session как явная альтернатива.",
    "Agent — текущий или рекомендованный агент с возможностью выбрать другого из доступных либо добавить нового.",
    "Model — текущая модель с заметным предложением выбрать другую для независимого второго мнения.",
    "Effort — уровень глубины и вычислительных усилий для проверки.",
    "Review instructions — необязательный фокус проверки, ограничения или дополнительные критерии.",
):
    add_bullet(doc, item)
add_body(doc, "В текущем чате те же параметры задаются непосредственно перед командой: scope корректируется attachments, Session остаётся текущей, Agent остаётся тем же, а Model, Effort и instructions можно изменить без отдельного диалога.")
add_body(doc, "Таким образом, в launch dialog можно выбрать другого агента и по умолчанию новую сессию, а в текущем чате агент и сессия фиксированы. Различие относится только к настройке запуска; после старта оба сценария получают одинаковый review lifecycle.")
add_body(doc, "Start Review доступна только при непустом scope, доступном агенте и выбранной модели. Подтверждение в dialog или команда в текущем чате создаёт одну review session с зафиксированным source context.")

add_step(doc, 3, "Пользователь запускает ревью и наблюдает за прогрессом")
add_body(doc, "При запуске через launch dialog IDE по умолчанию создаёт отдельную чат-сессию ревью и связывает её с исходным commit scope или diff. Если выбрана Current Session либо ревью запущено командой из текущего чата, в этом чате появляется отдельный review-блок с тем же lifecycle и теми же возможностями.")
add_body(doc, "В рабочий контекст входят выбранные изменения, attachments, комментарии текущего чата и комментарии, оставленные в файлах. Этот контекст сохраняется вместе с review session и используется в первом прогоне и последующих итерациях.")
add_body(doc, "Во время анализа progress component показывает все файлы scope и их состояния: Queued, Processing, Reviewed или Failed. Статусы остаются в компоненте и обновляются по мере обработки; когда обработан весь scope, компонент целиком исчезает и заменяется результатом ревью.")
add_body(doc, "Единственный способ остановить активную обработку — Stop в поле ввода чат-сессии. Отдельной кнопки остановки в progress component или другой поверхности нет.")

add_step(doc, 4, "Пользователь получает результат и выбирает глубину работы")
add_body(doc, "После обработки scope в чат-сессии появляется Review Preview со статусом Open. Оно показывает короткое и развёрнутое summary, число проверенных файлов и findings, основные риски и распределение по severity: Critical, Warning и Info.")
add_body(doc, "Каждый finding связан с файлом и anchor, содержит severity, описание, текущий статус и доступное исправление. Уже в Preview пользователь может ответить на finding, принять и применить исправление, отклонить его, оставить открытым или удалить; для severity-группы доступны применение всех подходящих исправлений и отклонение группы.")
add_body(doc, "Finding может находиться в состоянии Open, Accepted, Dismissed, Deleted или Pending update. Для V1 принятие finding и применение исправления являются одним состоянием Accepted. Изменения статусов сразу синхронизируются между Preview, Full Review, чат-сессией и связанным кодом и используются в counters и группировке по status.")
add_body(doc, "Preview, Full Review, review-блок чат-сессии и связанный код показывают одно состояние review. Findings, решения, counters, статусы и summary синхронизируются; переход между представлениями не создаёт новую проверку.")
add_body(doc, "С первого результата доступны три глобальных действия: Complete Review полностью и штатно завершает проверку; Submit Review продолжает её и позволяет перед запуском новой итерации написать сообщение агенту; Cancel Review отменяет ревью и предложенные в Preview изменения.")
add_info_note(doc, "Разграничение действий", "Cancel Review не используется вместо Complete Review. Complete фиксирует законченный результат, Submit продолжает работу, а Cancel означает явный отказ от ревью и его предложенных изменений.")
add_body(doc, "На этом этапе возникает единственная основная развилка:")
for item in (
    "Быстрое завершение — общей оценки достаточно, подробный разбор кода не нужен.",
    "Детальный разбор — нужно понять причины замечаний, их связь с кодом и при необходимости провести новую итерацию.",
):
    add_bullet(doc, item)
add_info_note(doc, "Особый случай", "Если findings нет, Preview показывает No findings и итоговый summary. Пользователь штатно завершает проверку через Complete Review. Submit Review недоступен без обратной связи, а Cancel Review остаётся отдельным действием только для явной отмены самого ревью и предложенных им изменений.")

add_step(doc, 5, "Сценарий 1 — быстро проверить и завершить")
add_info_note(doc, "Задача пользователя", "Получить независимую оценку изменений перед коммитом и принять решение без обязательного детального разбора каждого замечания.")

add_substep(doc, "5.1", "Оценить результат в Preview", ())
add_body(doc, "Пользователь читает короткое и развёрнутое summary, оценивает основные риски и распределение findings по severity, не покидая Preview.")
add_info_note(doc, "Подрезультат", "Пользователь понимает общий уровень риска и решает, достаточно ли полученной информации для завершения ревью.")

add_substep(doc, "5.2", "Выбрать, требуется ли обработка findings", ())
add_body(doc, "Если информации достаточно, пользователь не обязан обрабатывать findings и может сразу перейти к Complete Review. Если он хочет применить быстрое решение, выбирает масштаб действия:")
for item in (
    "Всё ревью — принять и применить все доступные исправления либо отклонить все предложения.",
    "Severity-группа — применить доступные исправления или отклонить все findings выбранной группы.",
    "Отдельный finding — ответить, принять и применить исправление, отклонить, оставить Open или удалить.",
):
    add_bullet(doc, item)
add_body(doc, "Групповое или массовое действие затрагивает только совместимые findings. Каждая карточка получает собственный статус, counters и оба summary пересчитываются; оставшиеся Open findings не блокируют завершение.")
add_info_note(doc, "Подрезультат", "Пользователь либо оставляет результат без изменений, либо применяет решение ко всему ревью, severity-группе или конкретному finding. Статусы, counters и summary отражают сделанный выбор.")

add_substep(doc, "5.3", "Выбрать дальнейшее действие со всем ревью", ())
add_body(doc, "После просмотра или быстрых решений пользователь выбирает один из трёх вариантов:")
for item in (
    "Complete Review — штатно завершить проверку. Статус меняется на Completed; итоговые counters и summary фиксируются, результат становится доступен только для чтения.",
    "Submit Review — написать сообщение агенту и продолжить проверку в новой итерации. В неё попадут ответы, решения, применённые исправления, Open и удалённые findings, комментарии из чата и файлов и актуальное состояние изменений.",
    "Cancel Review — явно отменить ревью и предложенные в Preview изменения. Статус меняется на Cancelled; это не альтернативный способ штатного завершения.",
):
    add_bullet(doc, item)
add_info_note(doc, "Подрезультат", "Complete фиксирует законченный результат; Submit создаёт продолжение текущего ревью; Cancel фиксирует явный отказ от ревью и отменяет предложенные изменения.")

add_substep(doc, "5.4", "Получить статус и результат выбранного действия", ())
for item in (
    "Complete Review — статус ревью меняется с Open на Completed. Итоговые counters и summary фиксируются, а review session становится доступна только для чтения.",
    "Cancel Review — статус ревью меняется с Open на Cancelled. Предложенные в Preview изменения отменяются, а review session становится доступна только для чтения.",
    "Submit Review — пользователь переходит в новую итерацию той же review session. Агент повторно обрабатывает актуальный scope и учитывает весь контекст предыдущих итераций. После обработки ревью получает промежуточный статус Updated; findings, counters и summary обновляются, после чего пользователь снова выбирает быструю или детальную работу с результатом.",
):
    add_bullet(doc, item)
add_body(doc, "Во время новой итерации пользователь снова видит progress и может остановить обработку только через Stop в поле ввода. После Complete Review или Cancel Review он возвращается к актуальным изменениям и решает, что делать дальше: выполнить commit, продолжить ручную работу или запустить новое Code Review.")
add_info_note(doc, "Подрезультат", "Complete приводит к финальному статусу Completed; Cancel — к финальному статусу Cancelled; Submit — к новой итерации и промежуточному статусу Updated после её обработки.")
add_info_note(doc, "Результат сценария", "Пользователь за короткое время получает независимую оценку и сам выбирает глубину реакции: просто прочитать результат, обработать весь результат, severity-группу или конкретный finding, а затем завершить, продолжить либо явно отменить ревью.")

add_step(doc, 6, "Сценарий 2 — детально разобрать и улучшить")
add_info_note(doc, "Задача пользователя", "Понять причины замечаний, связать их с полным контекстом кода и последовательно довести изменения до нужного качества.")

add_substep(doc, "6.1", "Открыть подробное представление", ())
add_body(doc, "Пользователь открывает результат в split view или Full Review. Это та же review session: все действия и статусы из Preview сохраняются и синхронно отображаются во всех представлениях.")
add_info_note(doc, "Подрезультат", "Пользователь получает больше рабочего пространства и сохраняет непрерывность контекста, решений и статусов текущего ревью.")

add_substep(doc, "6.2", "Настроить глубину и способ просмотра", ())
for item in (
    "Grouping — по severity, file, status или agent; filtering — по критичности и состоянию.",
    "Visual representation — от обзорного списка до представления с более тесной связью findings и кода.",
    "Code comparison — Split или Unified там, где нужно сравнение изменений.",
    "Full code context — переход из finding к полному файлу, окружающей реализации и зависимостям, а не только к изменённым строкам.",
    "Inline context — finding остаётся рядом с кодом, сохраняет severity и текущий статус.",
):
    add_bullet(doc, item)
add_info_note(doc, "Подрезультат", "Пользователь выбирает представление, которое даёт достаточную связь findings с изменениями и при необходимости открывает полный контекст файла и зависимостей.")

add_substep(doc, "6.3", "Разобрать findings и зафиксировать решения", ())
add_body(doc, "Пользователь отвечает на findings, запрашивает пояснения, принимает и применяет исправления, отклоняет рекомендации, оставляет их Open или удаляет. Действие можно применить на трёх уровнях: к отдельному finding, к severity-группе или ко всему результату.")
add_body(doc, "В Full Review пользователь также оставляет собственные комментарии к строкам и фрагментам кода. Ответы, решения, применённые исправления и ручные комментарии сохраняются как контекст текущего ревью.")
add_info_note(doc, "Подрезультат", "Каждый обработанный finding получает актуальный статус; массовые решения применяются только к совместимым findings; counters и summary пересчитываются, а накопленная обратная связь готова для следующей итерации.")

add_substep(doc, "6.4", "Выбрать дальнейшее действие со всем ревью", ())
add_body(doc, "После детального разбора пользователь выбирает один из трёх вариантов:")
for item in (
    "Complete Review — зафиксировать текущий результат как финальный. Статус меняется на Completed; counters и оба summary обновляются, результат становится доступен только для чтения.",
    "Submit Review — добавить сообщение агенту и передать накопленные решения, исправления и комментарии в следующую итерацию.",
    "Cancel Review — прекратить проверку, получить статус Cancelled и отменить предложенные в Preview изменения. Cancel не заменяет Complete Review.",
):
    add_bullet(doc, item)
add_info_note(doc, "Подрезультат", "Пользователь явно выбирает: зафиксировать результат, продолжить работу с накопленной обратной связью или отказаться от ревью и его предложенных изменений.")

add_substep(doc, "6.5", "Провести следующую итерацию, если выбран Submit Review", ())
add_body(doc, "IDE формирует единый feedback batch: сообщение агенту, ответы, принятые и отклонённые findings, применённые исправления, Open и удалённые items, комментарии из чата и файлов и актуальное состояние изменений.")
add_body(doc, "Статус меняется с Open на Updating. Агент обрабатывает актуальный scope внутри той же review session и учитывает полный контекст всех предыдущих итераций, а не только последнюю. Пользователь снова видит progress по файлам и может остановить обработку только кнопкой Stop в поле ввода.")
add_body(doc, "После обработки review получает статус Updated. IDE обновляет findings, их статусы, counters и оба summary. Пользователь снова оценивает новый результат: может быстро завершить ревью, перейти к подробному разбору, запустить ещё одну итерацию или отменить ревью.")
add_info_note(doc, "Подрезультат", "Пользователь получает обновлённый результат без потери решений, комментариев и истории предыдущих прогонов и может повторить цикл столько раз, сколько нужно.")

add_substep(doc, "6.6", "Завершить сценарий", ())
add_body(doc, "После Complete Review или Cancel Review дальнейшие действия с этой review session недоступны. Пользователь возвращается к актуальным изменениям и решает, что делать дальше: выполнить commit, продолжить ручную работу или запустить новое Code Review.")
add_info_note(doc, "Подрезультат", "Review получает финальный статус Completed или Cancelled, итоговые counters и summary обновляются, а Commit tool window показывает актуальные незакоммиченные изменения.")
add_info_note(doc, "Результат сценария", "Пользователь подробно разбирается в findings, связывает их с полным контекстом кода, передаёт агенту дополнительную информацию и проводит необходимое количество итераций до достижения нужного качества.")

add_step(doc, 7, "После финального состояния пользователь принимает решение об изменениях")
add_body(doc, "Commit tool window показывает текущее состояние незакоммиченных файлов с учётом применённых исправлений, последующих итераций и ручных изменений. Файлы сохраняют актуальные VCS-статусы: modified, added, deleted, renamed и untracked.")
add_body(doc, "Стандартный Commit tool window diff показывает выбранный файл относительно его базовой версии в VCS. Это не Review Preview: финальный статус и summary ревью остаются в связанной чат-сессии.")
add_body(doc, "Пользователь принимает следующее решение: выполнить commit, продолжить ручную работу или запустить новое Code Review. Новая проверка создаёт новую review session и не перезаписывает завершённую историю.")

h = doc.add_paragraph(style="Heading 2")
h.paragraph_format.space_before = Pt(8)
h.paragraph_format.space_after = Pt(4)
h.add_run("Источники")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0
r = p.add_run("Product overview: ")
set_run_font(r, size=9.5, color=INK, bold=True)
add_hyperlink(p, "AIUX-595 · Code Review overview", "https://urban-telegram-mvplegw.pages.github.io/AIUX-595/overview", size=9.5)
r = p.add_run("  ·  Interactive prototype: ")
set_run_font(r, size=9.5, color=INK, bold=True)
add_hyperlink(p, "AIUX-595 · Draft", "https://urban-telegram-mvplegw.pages.github.io/AIUX-595/draft", size=9.5)
r = p.add_run("  ·  ")
set_run_font(r, size=9.5, color=INK)
r = p.add_run("PRD: ")
set_run_font(r, size=9.5, color=INK, bold=True)
add_hyperlink(p, "PRD-857 · Code Review", "https://youtrack.jetbrains.com/issue/PRD-857", size=9.5)
r = p.add_run("  ·  Miro: ")
set_run_font(r, size=9.5, color=INK, bold=True)
add_hyperlink(p, "Code Review — user flow board", "https://miro.com/app/board/uXjVHN-Gcco=/?share_link_id=706548947908", size=9.5)
r = p.add_run("  ·  ")
set_run_font(r, size=9.5, color=INK)
r = p.add_run("Product decisions: ")
set_run_font(r, size=9.5, color=INK, bold=True)
r2 = p.add_run("комментарии пользователя к V1 flow, 11.08.2026.")
set_run_font(r2, size=9.5, color=INK)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "Code Review изменений в IDE — основной пользовательский flow V1"
doc.core_properties.subject = "GUI Code Review перед коммитом в IntelliJ IDEA"
doc.core_properties.author = "OpenAI"
doc.core_properties.keywords = "Code Review, IntelliJ IDEA, V1, user flow, GUI, pre-commit"
doc.save(OUTPUT)
print(OUTPUT.resolve())
