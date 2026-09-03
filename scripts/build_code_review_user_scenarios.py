from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("artifacts/Code_Review_V1_User_Scenarios_RU.docx")

# compact_reference_guide preset
INK = RGBColor(31, 45, 61)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(94, 104, 116)
PALE_BLUE = "EEF5FC"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_pr.rFonts.set(qn(f"w:{key}"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, cached, end))
    set_run_font(run, size=9, color=MUTED)


def create_numbering(doc, kind, start_at=1):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId"), "").isdigit()
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId"), "").isdigit()
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), str(start_at))
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if kind == "decimal" else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), "Calibri")
    r_pr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5" if kind == "decimal" else "1F4D78")
    r_pr.append(color)
    bold = OxmlElement("w:b")
    r_pr.append(bold)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_ref))
    p_pr.append(num_pr)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def add_body(doc, text, *, before=0, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_labeled_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = True
    lead = p.add_run(f"{label}. ")
    set_run_font(lead, size=11, color=DARK_BLUE, bold=True)
    rest = p.add_run(text)
    set_run_font(rest, size=11, color=INK)
    return p


def add_step(doc, num_id, text):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_bullet(doc, num_id, text):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.20
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run, size=10.7, color=INK)
    return p


def add_outcome(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.keep_together = True
    shade_paragraph(p, PALE_BLUE)
    label = p.add_run("РЕЗУЛЬТАТ  ")
    set_run_font(label, size=9, color=BLUE, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=10.8, color=INK)
    return p


doc = Document()
common_decimal_id = create_numbering(doc, "decimal")
quick_branch_decimal_id = create_numbering(doc, "decimal", start_at=7)
detailed_branch_decimal_id = create_numbering(doc, "decimal", start_at=7)
bullet_id = create_numbering(doc, "bullet")

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
header_p = header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_p.paragraph_format.space_after = Pt(0)
header_r = header_p.add_run("КОД-РЕВЬЮ · ЕДИНЫЙ ПОЛЬЗОВАТЕЛЬСКИЙ СЦЕНАРИЙ V1")
set_run_font(header_r, size=8.5, color=MUTED, bold=True)

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_p.paragraph_format.space_after = Pt(0)
footer_r = footer_p.add_run("Страница ")
set_run_font(footer_r, size=9, color=MUTED)
add_field(footer_p, "PAGE")

kicker = doc.add_paragraph()
kicker.paragraph_format.space_before = Pt(8)
kicker.paragraph_format.space_after = Pt(5)
kicker_r = kicker.add_run("PRODUCT SCENARIOS")
set_run_font(kicker_r, size=9, color=BLUE, bold=True)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(5)
title_r = title.add_run("Код-ревью: единый сценарий с развилкой")
set_run_font(title_r, size=24, color=INK, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
subtitle_r = subtitle.add_run("Общий путь запуска и две стратегии работы с результатом")
set_run_font(subtitle_r, size=12.5, color=MUTED)

problem_heading = doc.add_paragraph(style="Heading 2")
problem_heading.add_run("Проблема")
add_body(
    doc,
    "После завершения работы разработчику сложно быстро понять, готовы ли изменения к коммиту. Изменения могут быть распределены по нескольким файлам и затрагивать разные части системы, поэтому самостоятельная проверка не гарантирует, что весь объём просмотрен, а критичные проблемы не пропущены.",
)
add_body(
    doc,
    "При этом разным изменениям требуется разная глубина проверки: иногда достаточно быстро оценить основные риски, а иногда необходимо разобраться в каждом замечании и его полном контексте в коде. Ручной сбор этого контекста и повторная проверка после исправлений требуют дополнительного времени и внимания.",
    after=10,
)

lead = doc.add_paragraph()
lead.paragraph_format.space_after = Pt(12)
lead.paragraph_format.left_indent = Inches(0.12)
lead.paragraph_format.right_indent = Inches(0.08)
lead.paragraph_format.line_spacing = 1.20
lead.paragraph_format.keep_together = True
shade_paragraph(lead, PALE_BLUE)
lead_label = lead.add_run("ЦЕННОСТЬ ФИЧИ  ")
set_run_font(lead_label, size=9, color=BLUE, bold=True)
lead_text = lead.add_run(
    "Код-ревью помогает проверить подготовленные изменения, оценить риски и принять решения по результатам проверки до коммита — быстро или с подробным разбором и дополнительными итерациями."
)
set_run_font(lead_text, size=10.8, color=INK)

h1 = doc.add_paragraph(style="Heading 1")
h1.add_run("Единый пользовательский сценарий с развилкой")

add_labeled_paragraph(
    doc,
    "Ситуация",
    "Я закончил работу над задачей или очередную итерацию изменений и хочу проверить результат перед коммитом.",
)
add_labeled_paragraph(
    doc,
    "Цель",
    "Понять качество изменений, увидеть возможные проблемы и выбрать подходящую глубину работы с результатами Код-ревью.",
)

h2 = doc.add_paragraph(style="Heading 2")
h2.add_run("Общая последовательность")

for text in (
    "Я запускаю Код-ревью для подготовленных изменений.",
    "Определяю область проверки: весь набор изменений или только выбранную часть.",
    "Настраиваю проверку: выбираю агента, модель и формат продолжения работы, при необходимости добавляю дополнительные инструкции.",
    "Запускаю ревью. ИИ последовательно анализирует выбранные изменения, а я наблюдаю за прогрессом и понимаю, на каком этапе находится проверка. При необходимости я могу остановить процесс.",
    "После завершения получаю общую оценку изменений, краткое резюме и найденные проблемы, распределённые по критичности.",
    "Оцениваю полученный результат и выбираю дальнейший путь:",
):
    add_step(doc, common_decimal_id, text)

for text in (
    "Если мне достаточно общей оценки и я хочу быстро закончить проверку — перехожу к быстрому завершению ревью.",
    "Если мне нужно понять причины замечаний, изучить их связь с кодом и последовательно улучшить изменения — перехожу к детальному разбору.",
):
    add_bullet(doc, bullet_id, text)

h1 = doc.add_paragraph(style="Heading 1")
h1.add_run("Ветка 1. Быстро завершить ревью")

for text in (
    "Я читаю результат и определяю, нужны ли дополнительные действия.",
    "Если полученной информации достаточно, могу сразу завершить ревью, ничего не меняя и не обрабатывая замечания.",
    "Если я хочу обработать результат, выбираю подходящий уровень решения:",
):
    add_step(doc, quick_branch_decimal_id, text)

for text in (
    "Принять или отклонить все предложения ревью сразу.",
    "Принять или отклонить отдельную группу предложений.",
    "Принять или отклонить конкретное предложение.",
    "Оставить отдельные предложения без решения — это не мешает завершить ревью.",
):
    add_bullet(doc, bullet_id, text)

for text in (
    "Принятые исправления применяются к моим изменениям. Отклонённые предложения сохраняются как отклонённые, а итоговое состояние и резюме ревью обновляются.",
    "Я завершаю ревью. Его результат фиксируется как финальный, после чего дальнейшие действия с этой проверкой становятся недоступны.",
    "Возвращаюсь к актуальным изменениям и решаю, что делать дальше: выполнить коммит, продолжить работу или запустить новое ревью.",
):
    add_step(doc, quick_branch_decimal_id, text)

add_outcome(
    doc,
    "Я быстро получаю независимую оценку изменений и сам выбираю глубину реакции: просто прочитать результат, обработать всё ревью, отдельную группу или конкретное предложение.",
)

h1 = doc.add_paragraph(style="Heading 1")
h1.add_run("Ветка 2. Детально разобрать и улучшить изменения")

for text in (
    "Я перехожу к подробному разбору и выбираю глубину визуального представления результатов: от общей картины до тесной связи каждого замечания с кодом.",
    "При необходимости перехожу из конкретного замечания в полный контекст кода: изучаю связанный файл, окружающую реализацию и зависимости, чтобы понять причину проблемы и влияние предлагаемого изменения.",
    "Последовательно работаю с отдельными предложениями:",
):
    add_step(doc, detailed_branch_decimal_id, text)

for text in (
    "Запрашиваю пояснение.",
    "Оставляю собственный комментарий к предложению или фрагменту кода.",
    "Принимаю и применяю исправление.",
    "Отклоняю рекомендацию.",
    "Оставляю её без решения.",
    "Удаляю как неактуальную.",
):
    add_bullet(doc, bullet_id, text)

for text in (
    "Если нескольким предложениям подходит одинаковое решение, принимаю или отклоняю всю группу. Если одно решение подходит ко всему результату, принимаю или отклоняю все предложения сразу.",
    "Все мои ответы, комментарии, решения и применённые исправления сохраняются как часть контекста текущего ревью.",
    "После разбора выбираю один из трёх вариантов:",
):
    add_step(doc, detailed_branch_decimal_id, text)

for text in (
    "Продолжить ревью — передать накопленные решения, исправления и комментарии в следующую итерацию.",
    "Завершить ревью — зафиксировать текущий результат как финальный.",
    "Полностью отменить ревью — прекратить проверку без следующей итерации.",
):
    add_bullet(doc, bullet_id, text)

for text in (
    "Если я продолжаю ревью, ИИ анализирует актуальные изменения внутри той же проверки и учитывает полный контекст предыдущих итераций.",
    "После новой итерации получаю обновлённое резюме и новый набор результатов с учётом того, что изменилось. При необходимости повторяю цикл до достижения нужного качества.",
    "После завершения или полной отмены ревью получает финальное состояние, а его результат сохраняется как история проверки без дальнейших действий.",
    "Возвращаюсь к актуальным изменениям и решаю, что делать дальше: выполнить коммит, продолжить работу или запустить новое ревью.",
):
    add_step(doc, detailed_branch_decimal_id, text)

add_outcome(
    doc,
    "Я могу подробно разобраться в замечаниях, связать их с полным контекстом кода, дать ИИ дополнительную информацию и провести необходимое количество итераций до достижения нужного качества.",
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "Код-ревью — единый пользовательский сценарий с развилкой V1"
doc.core_properties.subject = "Общий путь Код-ревью и две ветки работы с результатом"
doc.core_properties.author = "OpenAI"
doc.core_properties.keywords = "Код-ревью, V1, пользовательские сценарии, code review"
doc.save(OUTPUT)
print(OUTPUT.resolve())
