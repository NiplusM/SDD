from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("artifacts/Code_Review_V1_User_Scenarios_EN.docx")

INK = RGBColor(30, 36, 43)
MUTED = RGBColor(91, 99, 109)
ACCENT = RGBColor(31, 77, 120)


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_pr.rFonts.set(qn(f"w:{key}"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, cached, end))
    set_run_font(run, size=9, color=MUTED)


def create_numbering(doc, kind, start_at=1):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId"), "").isdigit()
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId"), "").isdigit()
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), str(start_at))
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if kind == "decimal" else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "1080" if kind == "bullet" else "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "1080" if kind == "bullet" else "540")
    indent.set(qn("w:hanging"), "360" if kind == "bullet" else "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "90")
    spacing.set(qn("w:line"), "288")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), "Calibri")
    r_pr.append(fonts)
    if kind == "decimal":
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1F4D78")
        r_pr.append(color)
        bold = OxmlElement("w:b")
        r_pr.append(bold)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_ref))
    p_pr.append(num_pr)


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_labeled(doc, label, text, after=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.keep_together = True
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, size=11, color=INK, bold=True)
    rest = p.add_run(text)
    set_run_font(rest, size=11, color=INK)
    return p


def add_step(doc, num_id, text):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(4.5)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_bullet(doc, num_id, text):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run, size=10.8, color=INK)
    return p


def add_outcome(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.keep_together = True
    label = p.add_run("Outcome: ")
    set_run_font(label, size=11, color=INK, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=11, color=INK)
    return p


doc = Document()
common_decimal_id = create_numbering(doc, "decimal")
quick_branch_decimal_id = create_numbering(doc, "decimal")
detailed_branch_decimal_id = create_numbering(doc, "decimal")
bullet_id = create_numbering(doc, "bullet")

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.footer_distance = Inches(0.38)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

for name, size, before, after in (
    ("Heading 1", 16, 16, 8),
    ("Heading 2", 13, 12, 6),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = ACCENT
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_field(footer_p, "PAGE")

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
title_r = title.add_run("Code Review: One User Flow with Two Paths")
set_run_font(title_r, size=23, color=INK, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
subtitle_r = subtitle.add_run("A shared launch flow and two ways to work with the result")
set_run_font(subtitle_r, size=12, color=MUTED)

h2 = doc.add_paragraph(style="Heading 2")
h2.add_run("Problem")
add_body(
    doc,
    "After finishing their work, developers may struggle to quickly understand whether the changes are ready to commit. The changes may be spread across several files and affect different parts of the system, so a manual check does not guarantee that the full scope has been reviewed or that critical issues have not been missed.",
)
add_body(
    doc,
    "Different changes also require different levels of review. Sometimes a quick assessment of the main risks is enough; in other cases, every finding and its full code context must be investigated. Collecting this context manually and reviewing the changes again after fixes requires additional time and attention.",
    after=8,
)
add_labeled(
    doc,
    "Feature value",
    "Code Review helps users check prepared changes, assess risks, and make decisions based on the review results before committing—either quickly or through a detailed investigation and additional iterations.",
    after=12,
)

h1 = doc.add_paragraph(style="Heading 1")
h1.add_run("Shared User Flow with Two Paths")
add_labeled(
    doc,
    "Situation",
    "I have completed a task or another iteration of changes and want to review the result before committing it.",
)
add_labeled(
    doc,
    "Goal",
    "Understand the quality of the changes, identify potential issues, and choose the appropriate depth for working with the Code Review results.",
)

h2 = doc.add_paragraph(style="Heading 2")
h2.add_run("Shared sequence")
for text in (
    "I start Code Review for the prepared changes.",
    "I define the review scope: the complete change set or only a selected part.",
    "I configure the review by selecting an agent, model, and continuation format, and add optional instructions when needed.",
    "I start the review. AI analyzes the selected changes step by step while I monitor progress and understand which stage the review has reached. I can stop the process if necessary.",
    "When the review is complete, I receive an overall assessment of the changes, a short summary, and the identified issues grouped by severity.",
    "I evaluate the result and choose the path that matches my needs:",
):
    add_step(doc, common_decimal_id, text)

for text in (
    "If the high-level assessment is sufficient and I want to finish quickly, I continue with Path 1.",
    "If I need to understand the findings, connect them to the code, and improve the changes iteratively, I continue with Path 2.",
):
    add_bullet(doc, bullet_id, text)

h1 = doc.add_paragraph(style="Heading 1")
h1.add_run("Path 1. Complete the Review Quickly")

for text in (
    "I read the result and decide whether any additional action is needed.",
    "If the available information is sufficient, I can complete the review immediately without changing anything or processing the findings.",
    "If I want to process the result, I choose the appropriate level:",
):
    add_step(doc, quick_branch_decimal_id, text)

for text in (
    "Accept or dismiss all review proposals at once.",
    "Accept or dismiss a group of proposals.",
    "Accept or dismiss an individual proposal.",
    "Leave individual proposals unresolved; this does not prevent me from completing the review.",
):
    add_bullet(doc, bullet_id, text)

for text in (
    "Accepted fixes are applied to my changes. Dismissed proposals remain recorded as dismissed, while the final state and review summary are updated.",
    "I complete the review. Its result is recorded as final, after which no further actions are available for this review.",
    "I return to the current changes and decide whether to commit them, continue working, or start a new Code Review.",
):
    add_step(doc, quick_branch_decimal_id, text)

add_outcome(
    doc,
    "I quickly receive an independent assessment of the changes and choose how deeply to respond: simply read the result, process the entire review, a group, or an individual proposal.",
)

h1 = doc.add_paragraph(style="Heading 1")
h1.add_run("Path 2. Investigate in Detail and Improve the Changes")

for text in (
    "I open the detailed view and choose the depth of the visual representation, from a high-level overview to a close connection between each finding and the code.",
    "When necessary, I open the full code context from a specific finding and inspect the related file, surrounding implementation, and dependencies to understand the cause of the issue and the impact of the proposed change.",
    "I work through individual proposals one by one:",
):
    add_step(doc, detailed_branch_decimal_id, text)

for text in (
    "Request clarification.",
    "Add my own comment to a proposal or code fragment.",
    "Accept and apply the fix.",
    "Dismiss the recommendation.",
    "Leave it unresolved.",
    "Delete it as no longer relevant.",
):
    add_bullet(doc, bullet_id, text)

for text in (
    "If the same decision applies to several proposals, I accept or dismiss the entire group. If one decision applies to the whole result, I accept or dismiss all proposals at once.",
    "All of my replies, comments, decisions, and applied fixes are retained as part of the current review context.",
    "After the investigation, I choose one of three options:",
):
    add_step(doc, detailed_branch_decimal_id, text)

for text in (
    "Continue the review—pass the accumulated decisions, fixes, and comments to the next iteration.",
    "Complete the review—record the current result as final.",
    "Cancel the review completely—stop the review without another iteration.",
):
    add_bullet(doc, bullet_id, text)

for text in (
    "If I continue the review, AI analyzes the current changes within the same review and considers the full context of all previous iterations.",
    "After the new iteration, I receive an updated summary and a new set of results reflecting what has changed. I can repeat the cycle as needed until the required quality is reached.",
    "After the review is completed or fully cancelled, it receives a final state and its result is retained as review history without further actions.",
    "I return to the current changes and decide whether to commit them, continue working, or start a new Code Review.",
):
    add_step(doc, detailed_branch_decimal_id, text)

add_outcome(
    doc,
    "I can investigate findings in detail, connect them to the full code context, provide the agent with additional information, and run as many iterations as needed to reach the required quality.",
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "Code Review — One User Flow with Two Paths"
doc.core_properties.subject = "A shared Code Review flow and two ways to work with the result"
doc.core_properties.author = "OpenAI"
doc.core_properties.keywords = "Code Review, V1, user scenarios, review workflow"
doc.save(OUTPUT)
print(OUTPUT.resolve())
